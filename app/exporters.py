from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
import os
import re

class WordExporter:
    """Fixed Word exporter"""
    
    def __init__(self):
        pass
    
    def create_document_from_text(self, text, output_path, title=None):
        """
        Create Word document from plain text
        """
        try:
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
            
            # Add title
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            
            # Add content
            if not text or not isinstance(text, str):
                text = "No content available"
            
            paragraphs = text.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    
                    # Format
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            
            doc.save(output_path)
            print(f"Word document saved: {output_path}")
            return output_path
        
        except Exception as e:
            print(f"Word export error: {str(e)}")
            raise Exception(f"Failed to create Word document: {str(e)}")
    
    def create_document_from_structure(self, structure, output_path, title=None):
        """
        Create Word with structure - FIXED
        """
        try:
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)
            
            # Add title
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            
            # Check if structure is valid
            if not structure or not isinstance(structure, dict):
                print("Invalid structure, using plain text")
                text = structure if isinstance(structure, str) else "No content"
                return self.create_document_from_text(text, output_path, title)
            
            # Get structure elements
            structure_data = structure.get('structure', {})
            structure_elements = structure_data.get('structure', []) if isinstance(structure_data, dict) else []
            
            if not structure_elements:
                # No structure, use text
                text = structure.get('text', 'No content')
                return self.create_document_from_text(text, output_path, title)
            
            # Process structured elements
            for element in structure_elements:
                try:
                    if not isinstance(element, dict):
                        continue
                    
                    elem_type = element.get('type', '')
                    
                    if elem_type == 'heading':
                        content = element.get('content', '')
                        level = element.get('level', 2)
                        if content:
                            doc.add_heading(str(content), min(level, 9))
                    
                    elif elem_type == 'paragraph':
                        content = element.get('content', '')
                        if content:
                            para = doc.add_paragraph(str(content))
                            for run in para.runs:
                                run.font.name = 'Calibri'
                                run.font.size = Pt(11)
                    
                    elif elem_type == 'list':
                        items = element.get('items', [])
                        list_type = element.get('list_type', 'bullet')
                        style_name = 'List Number' if list_type == 'numbered' else 'List Bullet'
                        
                        for item in items:
                            if item:
                                clean_item = re.sub(r'^[-•*\d.)\s]+', '', str(item)).strip()
                                if clean_item:
                                    try:
                                        doc.add_paragraph(clean_item, style=style_name)
                                    except:
                                        # Fallback if style doesn't exist
                                        doc.add_paragraph(clean_item, style='List Paragraph')
                
                except Exception as e:
                    print(f"Error processing element: {e}")
                    continue
            
            # Add detected tables if available
            tables = structure.get('tables') if isinstance(structure, dict) else None
            # Handle case where tables might be in content dict passed as 'structure'
            if not tables and isinstance(structure, dict) and 'tables' not in structure and structure.get('structure') is None:
                 # This happens if structure is the content dict
                 tables = structure.get('tables')

            if tables and len(tables) > 0:
                for i, table in enumerate(tables):
                    table_data = table.get('data') if isinstance(table, dict) else table
                    if table_data:
                        self.add_table_to_document(doc, table_data)
                        doc.add_paragraph() # Spacing
            
            doc.save(output_path)
            print(f"Structured Word document saved: {output_path}")
            return output_path
        
        except Exception as e:
            print(f"Structured Word export error: {str(e)}")
            # Fallback to simple text export
            try:
                text = structure.get('text', 'No content') if isinstance(structure, dict) else str(structure)
                return self.create_document_from_text(text, output_path, title)
            except:
                raise Exception(f"Failed to create Word document: {str(e)}")

    def add_table_to_document(self, doc, table_data):
        """
        Add a table to Word document
        """
        if not table_data or len(table_data) == 0:
            return
        
        try:
            # Normalize table data
            max_cols = max(len(row) for row in table_data)
            rows = len(table_data)
            
            table = doc.add_table(rows=rows, cols=max_cols)
            table.style = 'Table Grid'
            
            for i, row in enumerate(table_data):
                for j, cell_text in enumerate(row):
                    if j < max_cols:
                        table.cell(i, j).text = str(cell_text)
        except Exception as e:
            print(f"Error adding table to Word doc: {e}")

class ExcelExporter:
    """Fixed Excel exporter"""
    
    def __init__(self):
        pass
    
    def create_excel_from_tables(self, tables, output_path, sheet_names=None):
        """
        Create Excel from tables
        """
        try:
            if not tables:
                raise Exception("No tables to export")
            
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                for i, table in enumerate(tables):
                    if not table or len(table) == 0:
                        continue
                    
                    try:
                        # Normalize row lengths
                        max_cols = 0
                        if len(table) > 0:
                            max_cols = max(len(row) for row in table)
                        
                        if max_cols == 0:
                            continue
                            
                        # Pad rows
                        normalized_table = []
                        for row in table:
                            if len(row) < max_cols:
                                row = row + [''] * (max_cols - len(row))
                            normalized_table.append(row[:max_cols])
                        
                        # Prepare data
                        if len(normalized_table) > 1:
                            headers = normalized_table[0]
                            data_rows = normalized_table[1:]
                        else:
                            headers = [f"Column_{j+1}" for j in range(max_cols)]
                            data_rows = normalized_table
                        
                        # Create DataFrame
                        df = pd.DataFrame(data_rows, columns=headers)
                        
                        # Clean
                        df = df.dropna(how='all')
                        
                        # Sheet name
                        if sheet_names and i < len(sheet_names):
                            sheet_name = self.sanitize_sheet_name(sheet_names[i])
                        else:
                            sheet_name = f'Table_{i+1}'
                        
                        # Write
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                        
                        # Style
                        worksheet = writer.sheets[sheet_name]
                        self._style_worksheet(worksheet, df)
                        
                    except Exception as e:
                        print(f"Error processing table {i}: {e}")
                        continue
            
            print(f"Excel file saved: {output_path}")
            return output_path
        
        except Exception as e:
            print(f"Excel export error: {str(e)}")
            raise Exception(f"Failed to create Excel: {str(e)}")
    
    def sanitize_sheet_name(self, name):
        """Create valid sheet name"""
        name = re.sub(r'[\\/*?\[\]:]', '_', str(name))
        name = name[:31]
        return name if name else 'Sheet'
    
    def _style_worksheet(self, worksheet, df):
        """Apply styling"""
        try:
            # Header style
            header_fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
            header_font = Font(color='FFFFFF', bold=True, size=11)
            
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Auto-adjust columns
            for column in worksheet.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                
                for cell in column:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                
                adjusted_width = min(max(max_length + 2, 10), 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Freeze header
            worksheet.freeze_panes = 'A2'
            
        except Exception as e:
            print(f"Styling error: {e}")

class MixedExporter:
    """Export both formats"""
    
    def __init__(self):
        self.word_exporter = WordExporter()
        self.excel_exporter = ExcelExporter()
    
    def create_mixed_document(self, content, output_dir, base_name='document'):
        """
        Create Word and Excel
        """
        try:
            results = {}
            
            # Word document
            if content.get('text') or content.get('structure'):
                word_path = os.path.join(output_dir, f'{base_name}.docx')
                
                try:
                    structure = content.get('structure')
                    if structure and isinstance(structure, dict) and structure.get('structure'):
                        self.word_exporter.create_document_from_structure(
                            structure,
                            word_path,
                            title=base_name.replace('_', ' ').title()
                        )
                    else:
                        text = content.get('text', 'No content')
                        self.word_exporter.create_document_from_text(
                            text,
                            word_path,
                            title=base_name.replace('_', ' ').title()
                        )
                    
                    results['word'] = word_path
                    
                except Exception as e:
                    print(f"Word export failed: {e}")
            
            # Excel document
            if content.get('tables'):
                tables = content['tables']
                if tables:
                    excel_path = os.path.join(output_dir, f'{base_name}_tables.xlsx')
                    
                    try:
                        table_data = []
                        for table in tables:
                            if isinstance(table, dict) and table.get('data'):
                                table_data.append(table['data'])
                            elif isinstance(table, list):
                                table_data.append(table)
                        
                        if table_data:
                            self.excel_exporter.create_excel_from_tables(table_data, excel_path)
                            results['excel'] = excel_path
                    
                    except Exception as e:
                        print(f"Excel export failed: {e}")
            
            return results
        
        except Exception as e:
            print(f"Mixed export error: {str(e)}")
            raise Exception(f"Failed to create documents: {str(e)}")